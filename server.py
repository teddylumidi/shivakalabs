#!/usr/bin/env python3
import os
import re
import secrets
import gzip
import io
from datetime import datetime, timedelta
from flask import Flask, request, jsonify, send_file, session, make_response
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_talisman import Talisman
from flask_wtf.csrf import CSRFProtect
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
import requests
import html
import bleach

app = Flask(__name__)
app.config['COMPRESS_LEVEL'] = 9
app.config['JSON_SORT_KEYS'] = False

# Security Configuration
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', secrets.token_hex(32))
app.config['SESSION_COOKIE_SECURE'] = True
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=1)
app.config['WTF_CSRF_TIME_LIMIT'] = None

# Compression middleware
@app.before_request
def compress_request():
    request.is_gzip = 'gzip' in request.headers.get('Accept-Encoding', '')

@app.after_request
def compress_response(response):
    if getattr(request, 'is_gzip', False) and response.status_code == 200:
        try:
            if response.direct_passthrough:
                return response
            data = response.get_data()
            if len(data) > 500:
                response.data = gzip.compress(data)
                response.headers['Content-Encoding'] = 'gzip'
                response.headers['Content-Length'] = len(response.data)
        except:
            pass
    return response

# Security & Performance Headers
@app.after_request
def set_security_headers(response):
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    response.headers['Permissions-Policy'] = 'geolocation=*, microphone=(), camera=()'
    response.headers['Strict-Transport-Security'] = 'max-age=31536000'
    
    # Aggressive caching for static assets
    if request.path.endswith(('.js', '.css', '.png', '.jpg', '.jpeg', '.gif', '.svg', '.woff', '.woff2', '.ttf')):
        response.headers['Cache-Control'] = 'public, max-age=31536000, immutable'
        response.headers['ETag'] = None
    elif request.path.endswith('.html'):
        response.headers['Cache-Control'] = 'public, max-age=3600, must-revalidate'
    else:
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
    
    return response

# Initialize security tools
CORS(app, resources={r"/api/*": {"origins": os.getenv('ALLOWED_ORIGINS', '*')}})
Talisman(app, 
    force_https=False,
    strict_transport_security=True,
    strict_transport_security_max_age=31536000,
    content_security_policy={
        'default-src': "'self'",
        'script-src': ["'self'", "'unsafe-inline'", "https://cdn.tailwindcss.com", "https://js.paystack.co"],
        'style-src': ["'self'", "'unsafe-inline'", "https://cdn.tailwindcss.com"],
        'img-src': ["'self'", "data:", "https:", "blob:"],
        'connect-src': ["'self'", "https://api.paystack.co"],
        'frame-src': ["'self'", "https://checkout.paystack.com"],
        'base-uri': ["'self'"],
        'form-action': ["'self'"]
    }
)

csrf = CSRFProtect(app)
limiter = Limiter(app=app, key_func=get_remote_address, default_limits=["200 per day", "50 per hour"])

# Paystack API Keys
PAYSTACK_SECRET_KEY = os.getenv('PAYSTACK_SECRET_KEY', '238b292d4da31544')
PAYSTACK_PUBLIC_KEY = os.getenv('PAYSTACK_PUBLIC_KEY', 'pk_test_35cd142d3174518a92af53c22e98d20b297c4aca')

# Security utilities
def sanitize_input(data):
    """Sanitize user input to prevent XSS"""
    if isinstance(data, dict):
        return {k: sanitize_input(v) for k, v in data.items()}
    if isinstance(data, list):
        return [sanitize_input(item) for item in data]
    if isinstance(data, str):
        data = html.escape(data)
        allowed_tags = []
        allowed_attributes = {}
        return bleach.clean(data, tags=allowed_tags, attributes=allowed_attributes, strip=True)
    return data

def validate_email(email):
    """Validate email format"""
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email) is not None

def validate_phone(phone):
    """Validate phone format (Kenya or international)"""
    phone = str(phone).replace(' ', '')
    # Kenya: 254XXXXXXXXX, +254XXXXXXXXX, 07XXXXXXXXX
    # International: +1-999 with 7-15 digits
    pattern = r'^254\d{9}$|^\+254\d{9}$|^0\d{9}$|^\+\d{7,15}$'
    return re.match(pattern, phone) is not None

def generate_csrf_token():
    """Generate CSRF token"""
    return secrets.token_urlsafe(32)

@app.before_request
def before_request():
    """Security checks before each request"""
    if request.method == 'POST':
        if not request.is_json and request.method == 'POST':
            return jsonify({'error': 'Invalid content type'}), 400
        
        data = request.get_json(silent=True)
        if data is None and request.method == 'POST':
            return jsonify({'error': 'Invalid JSON'}), 400

@app.after_request
def after_request(response):
    """Security headers after each request"""
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    response.headers['Permissions-Policy'] = 'geolocation=(), microphone=(), camera=()'
    return response

# Routes
@app.route('/')
def serve_index():
    """Serve main index"""
    response = make_response(send_file('index.html'))
    response.headers['Cache-Control'] = 'public, max-age=3600'
    response.headers['Content-Type'] = 'text/html; charset=utf-8'
    return response

@app.route('/<path:path>')
def serve_static(path):
    """Serve static files with caching"""
    try:
        if path.endswith(('.js', '.css', '.png', '.jpg', '.jpeg', '.gif', '.svg', '.woff', '.woff2')):
            response = make_response(send_file(path))
            response.headers['Cache-Control'] = 'public, max-age=31536000'
            return response
        return send_file(path)
    except:
        return send_file('index.html')

@app.route('/api/csrf-token', methods=['GET'])
@limiter.limit("100 per hour")
def get_csrf_token():
    """Get CSRF token"""
    token = generate_csrf_token()
    session['csrf_token'] = token
    response = make_response(jsonify({'token': token}))
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    return response

@app.route('/api/process-payment', methods=['POST'])
@limiter.limit("10 per hour")
@csrf.exempt
def process_payment():
    """Process global payment (Paystack or Stripe)"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        data = sanitize_input(data)
        reference = data.get('reference', '').strip()
        email = data.get('email', '').strip()
        phone = data.get('phone', '').strip()
        gateway = data.get('gateway', 'paystack').strip()
        
        if not all([reference, email, phone]):
            return jsonify({'success': False, 'error': 'Missing required fields'}), 400
        
        if not validate_email(email):
            return jsonify({'success': False, 'error': 'Invalid email'}), 400
        
        if not validate_phone(phone):
            return jsonify({'success': False, 'error': 'Invalid phone'}), 400
        
        if not re.match(r'^[a-zA-Z0-9_-]+$', reference):
            return jsonify({'success': False, 'error': 'Invalid reference'}), 400
        
        if gateway == 'paystack':
            headers = {'Authorization': f'Bearer {PAYSTACK_SECRET_KEY}', 'Content-Type': 'application/json'}
            response = requests.get(f'https://api.paystack.co/transaction/verify/{reference}', headers=headers, timeout=10)
            if response.status_code == 200 and response.json().get('data', {}).get('status') == 'success':
                return jsonify({'success': True, 'message': 'Payment verified'}), 200
        
        return jsonify({'success': False, 'error': 'Payment verification failed'}), 401
    
    except requests.exceptions.RequestException:
        return jsonify({'success': False, 'error': 'Payment service error'}), 503
    except Exception as e:
        app.logger.error(f'Payment error: {str(e)}')
        return jsonify({'success': False, 'error': 'Processing error'}), 500

@app.route('/api/initiate-payment', methods=['POST'])
@limiter.limit("10 per hour")
@csrf.exempt
def initiate_payment():
    """Initiate payment for international customers"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        data = sanitize_input(data)
        email = data.get('email', '').strip()
        amount = data.get('amount', 0)
        currency = data.get('currency', 'USD').strip()
        
        if not validate_email(email) or amount <= 0:
            return jsonify({'success': False, 'error': 'Invalid data'}), 400
        
        # For now, return success - integrate Stripe when keys available
        return jsonify({
            'success': True,
            'checkout_url': f'https://checkout.stripe.com?amount={amount}&currency={currency}&email={email}'
        }), 200
    
    except Exception as e:
        app.logger.error(f'Payment initiation error: {str(e)}')
        return jsonify({'success': False, 'error': 'Initiation failed'}), 500

@app.route('/api/generate-document', methods=['POST'])
@limiter.limit("10 per hour")
@csrf.exempt
def generate_document():
    """Generate CV/Cover Letter with validation"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
        
        # Sanitize all input
        data = sanitize_input(data)
        
        package_type = data.get('packageType', '').strip()
        work_exp = data.get('work_experience', '').strip()
        education = data.get('education', '').strip()
        skills = data.get('skills', '').strip()
        
        # Validate
        if package_type not in ['cv', 'cover', 'both']:
            return jsonify({'success': False, 'error': 'Invalid package type'}), 400
        
        if not all([work_exp, education, skills]):
            return jsonify({'success': False, 'error': 'Missing required fields'}), 400
        
        if len(work_exp) > 5000 or len(education) > 5000 or len(skills) > 2000:
            return jsonify({'success': False, 'error': 'Input too long'}), 400
        
        # Generate documents
        if package_type in ['cv', 'both']:
            cv_filename = f"cv_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            generate_cv_pdf(cv_filename, work_exp, education, skills)
        
        if package_type in ['cover', 'both']:
            letter_filename = f"cover_letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            generate_cover_letter_docx(letter_filename)
        
        return jsonify({'success': True, 'message': 'Documents generated'}), 200
    
    except Exception as e:
        app.logger.error(f'Document generation error: {str(e)}')
        return jsonify({'success': False, 'error': 'Generation failed'}), 500

def generate_cv_pdf(filename, work_exp, education, skills):
    """Generate professional CV"""
    doc = SimpleDocTemplate(filename, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#D4AF37',
        spaceAfter=6,
        alignment=1
    )
    
    story.append(Paragraph("Professional CV", title_style))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph("Work Experience", styles['Heading2']))
    story.append(Paragraph(work_exp, styles['BodyText']))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph("Education", styles['Heading2']))
    story.append(Paragraph(education, styles['BodyText']))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph("Skills", styles['Heading2']))
    story.append(Paragraph(skills, styles['BodyText']))
    
    doc.build(story)

def generate_cover_letter_docx(filename):
    """Generate professional cover letter"""
    doc = Document()
    doc.add_heading('Professional Cover Letter', 0)
    doc.add_paragraph()
    doc.add_paragraph('Dear Hiring Manager,')
    doc.add_paragraph()
    doc.add_paragraph('I am writing to express my strong interest in the position.')
    doc.add_paragraph()
    doc.add_paragraph('Best regards,')
    doc.save(filename)

@app.errorhandler(429)
def ratelimit_handler(e):
    """Rate limit handler"""
    return jsonify({'error': 'Rate limit exceeded', 'retry_after': 60}), 429

@app.errorhandler(404)
def not_found(e):
    """404 handler"""
    return send_file('index.html'), 200

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=False)
